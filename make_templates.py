from docx import Document
from docx.shared import Pt

def build_rental():
    doc = Document()
    doc.add_heading('RENTAL AGREEMENT', level=1)
    p = doc.add_paragraph()
    p.add_run(
        "This Rental Agreement is made on {{start_date}} between "
        "{{landlord_name}} (\"Landlord\") and {{tenant_name}} (\"Tenant\") "
        "for the residential property located at {{property_address}}."
    )
    doc.add_paragraph("Monthly Rent: Rs. {{monthly_rent}}")
    doc.add_paragraph("Security Deposit: Rs. {{security_deposit}}")
    doc.add_paragraph("Agreement Duration: {{agreement_duration_months}} months")
    doc.add_paragraph(
        "Both parties agree to the terms and conditions of this rental "
        "agreement as per applicable Kerala tenancy regulations."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Agreement Number: {{agreement_number}}")
    doc.add_paragraph("Generated on: {{today_date}}")
    doc.add_paragraph("Stamp Duty Applicable: Rs. {{stamp_duty_amount}}")
    doc.add_paragraph("")
    watermark_p = doc.add_paragraph()
    run = watermark_p.add_run("{{watermark_text}}")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph("")
    doc.add_paragraph("_____________________          _____________________")
    doc.add_paragraph("Landlord Signature                    Tenant Signature")
    doc.save("templates_docx/rental_agreement.docx")

def build_shop():
    doc = Document()
    doc.add_heading('SHOP / COMMERCIAL LEASE AGREEMENT', level=1)
    p = doc.add_paragraph()
    p.add_run(
        "This Commercial Lease Agreement is made on {{start_date}} between "
        "{{landlord_name}} (\"Landlord\") and {{tenant_name}} (\"Tenant\") "
        "for the commercial property located at {{shop_address}}, "
        "to be used for the business of {{business_type}}."
    )
    doc.add_paragraph("Monthly Rent: Rs. {{monthly_rent}}")
    doc.add_paragraph("Security Deposit: Rs. {{security_deposit}}")
    doc.add_paragraph("Agreement Duration: {{agreement_duration_months}} months")
    doc.add_paragraph(
        "Both parties agree to the terms and conditions of this commercial "
        "lease as per applicable Kerala regulations."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Agreement Number: {{agreement_number}}")
    doc.add_paragraph("Generated on: {{today_date}}")
    doc.add_paragraph("Stamp Duty Applicable: Rs. {{stamp_duty_amount}}")
    doc.add_paragraph("")
    watermark_p = doc.add_paragraph()
    run = watermark_p.add_run("{{watermark_text}}")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph("")
    doc.add_paragraph("_____________________          _____________________")
    doc.add_paragraph("Landlord Signature                    Tenant Signature")
    doc.save("templates_docx/shop_agreement.docx")

build_rental()
build_shop()
print("templates created")
