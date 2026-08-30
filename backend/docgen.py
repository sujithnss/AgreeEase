"""
Fills the appropriate .docx template with verified data using docxtpl.
Same template is reused for both the customer-facing watermarked draft
and the final in-house print copy — only `watermark_text` differs.

Stamp duty calculation here is a placeholder formula — replace with the
actual Kerala stamp duty rules the business uses.
"""

import os
import re
import shutil
import subprocess
import datetime
from xml.sax.saxutils import quoteattr
from docxtpl import DocxTemplate
from docx.oxml import parse_xml
from templates_config import get_template, get_template_file, duration_months_for
from dateutils import calculate_agreement_end_date
from malayalam_numwords import (
    number_to_malayalam_words,
    ordinal_malayalam_words,
    month_name_malayalam,
)

# Resolve paths relative to the project root (one level up from backend/),
# so this works regardless of the directory the app is launched from.
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _slugify(value: str, max_len: int = 30) -> str:
    """Turns a name/phone into a filesystem- and URL-safe slug, e.g.
    'Rahul Menon' -> 'rahul-menon'. Used to tag filenames so a document
    can be traced back to its customer just by its filename."""
    if not value:
        return "unknown"
    value = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return value[:max_len] or "unknown"


def _add_diagonal_watermark(document, text: str) -> None:
    """Injects a real diagonal, translucent watermark into every section's
    header — the same VML shape trick Word's own Insert > Watermark > Text
    feature generates, so LibreOffice (which already has to render legacy
    Word watermarks correctly) converts it the same way. python-docx has no
    high-level API for this, so it's added as raw OOXML on the header's
    paragraph run rather than through the normal python-docx object model.
    """
    watermark_xml = f"""
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office">
      <v:shapetype id="_x0000_t136" coordsize="1600,21600" o:spt="136"
          adj="10800" path="m@7,0l@8,0m@5,21600l@6,21600e">
        <v:formulas>
          <v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/>
          <v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/>
          <v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 @1"/>
          <v:f eqn="if @0 @4 @2"/><v:f eqn="if @0 #0 21600"/>
          <v:f eqn="if @0 0 #0"/><v:f eqn="if @0 21600 @8"/>
          <v:f eqn="if @0 @9 0"/><v:f eqn="if @0 @6 @5"/>
        </v:formulas>
        <v:path textpathok="t" o:connecttype="custom"
            o:connectlocs="10800,0;0,10800;10800,21600;21600,10800"
            o:connectangles="270,180,90,0"/>
        <v:textpath on="t" fitshape="t"/>
        <v:handles><v:h position="#0,bottomRight" xrange="0,21600"/></v:handles>
      </v:shapetype>
      <v:shape id="AgreeEaseWatermark" o:spid="_x0000_s2050" type="#_x0000_t136"
          style="position:absolute;margin-left:0;margin-top:0;width:415pt;height:207.5pt;
          rotation:315;z-index:-251654144;mso-position-horizontal:center;
          mso-position-horizontal-relative:margin;mso-position-vertical:center;
          mso-position-vertical-relative:margin"
          o:allowincell="f" fillcolor="silver" stroked="f">
        <v:fill opacity=".5"/>
        <v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt" string={quoteattr(text)}/>
      </v:shape>
    </w:pict>
    """
    for section in document.sections:
        header = section.header
        # Sections default to "linked to previous" with no real header part
        # of their own — paragraphs added before flipping this would be
        # written to a throwaway element and never actually saved.
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = paragraph.add_run()
        run._r.append(parse_xml(watermark_xml))


def calculate_stamp_duty(agreement_type: str, fields: dict) -> str:
    """
    PLACEHOLDER LOGIC — replace with real Kerala stamp duty rules.
    Rental/lease stamp duty is often a percentage of (annual rent + deposit),
    but exact rates depend on agreement type, duration, and local rules.
    Keep this as a single function so it's easy for the business's team
    to correct without touching the rest of the code.
    """
    try:
        monthly_rent = float(str(fields.get("monthly_rent", 0)).replace(",", ""))
        deposit = float(str(fields.get("security_deposit", 0)).replace(",", ""))
        months = int(fields.get("agreement_duration_months", 11))
    except (ValueError, TypeError):
        return "TO BE CALCULATED BY STAFF"

    annual_value = (monthly_rent * 12) + deposit
    # Placeholder: 1% of annual value, staff should verify against actual rules
    duty = round(annual_value * 0.01, 2)
    return f"{duty:,.2f}"


def generate_document(
    request_id: int,
    agreement_type: str,
    fields: dict,
    draft: bool = True,
    customer_phone: str = "",
    customer_name: str = "",
    language: str = "malayalam",
) -> str:
    """
    Generates the filled document.
    draft=True  -> watermarked "DRAFT - NOT VALID" copy for customer review
    draft=False -> clean copy for in-house printing only (staff use)
    language    -> which template file to use ("malayalam" or "english");
                   most agreements are drafted in Malayalam in practice, so
                   that's the default. Falls back to whatever's available
                   for agreement types without a Malayalam template yet.

    The output filename is tagged with the customer's phone number and
    name (e.g. request_12_919876543210_rahul-menon_final.docx) so a
    document can be identified and traced back to its customer directly
    from the generated/ folder, in addition to the DB record.
    """
    template_info = get_template(agreement_type)
    if not template_info:
        raise ValueError(f"Unknown agreement type: {agreement_type}")

    template_path = os.path.join(PROJECT_ROOT, get_template_file(agreement_type, language))
    doc = DocxTemplate(template_path)

    context = dict(fields)
    context["agreement_number"] = f"AGR-{request_id:06d}"
    today = datetime.date.today()
    context["today_date"] = today.strftime("%d-%m-%Y")
    context["stamp_duty_amount"] = calculate_stamp_duty(agreement_type, fields)
    # No longer a body-text banner (see _add_diagonal_watermark below) —
    # kept as an empty string rather than removed so existing templates
    # with a {{ watermark_text }} placeholder still render (as a blank line).
    context["watermark_text"] = ""
    # Resolves to whatever the customer/staff supplied, falling back to
    # the template's default (see FIXED_DURATION_MONTHS) -- covers both
    # older requests made before this was a collected field and the case
    # where the customer never stated a duration. Both templates render
    # this via {{ agreement_duration_months }} rather than a fixed number,
    # so a customer asking for 24 or 36 months instead of the traditional
    # 11 doesn't need a different template.
    context["agreement_duration_months"] = duration_months_for(agreement_type, fields)
    end_date = calculate_agreement_end_date(context)
    context["end_date"] = end_date.strftime("%d-%m-%Y") if end_date else ""

    # Malayalam word-form context, for templates (like rental_agreement_ml)
    # that follow the vendor's convention of spelling dates and amounts in
    # words alongside digits — see malayalam_numwords.py. Harmless no-ops
    # for templates that don't reference these placeholders.
    context["today_year"] = today.year
    context["today_year_words"] = number_to_malayalam_words(today.year)
    context["today_month_name_ml"] = month_name_malayalam(today.month)
    context["today_day_ordinal_ml"] = ordinal_malayalam_words(today.day)
    try:
        context["monthly_rent_words"] = number_to_malayalam_words(
            int(float(str(fields.get("monthly_rent", 0)).replace(",", "")))
        )
    except (ValueError, TypeError):
        context["monthly_rent_words"] = ""
    try:
        context["security_deposit_words"] = number_to_malayalam_words(
            int(float(str(fields.get("security_deposit", 0)).replace(",", "")))
        )
    except (ValueError, TypeError):
        context["security_deposit_words"] = ""
    try:
        context["agreement_duration_months_words_ml"] = number_to_malayalam_words(
            int(context["agreement_duration_months"])
        )
    except (ValueError, TypeError):
        context["agreement_duration_months_words_ml"] = ""

    doc.render(context)

    if draft:
        # NOT doc.get_docx() -- that reloads a fresh unrendered copy of the
        # template whenever is_rendered is already True (which it is, right
        # after doc.render() above), silently discarding everything render()
        # just filled in. doc.docx is the actual rendered Document object.
        _add_diagonal_watermark(doc.docx, "DRAFT - NOT VALID FOR STAMPING")

    suffix = "draft" if draft else "final"
    lang_tag = "ml" if language == "malayalam" else "en"
    phone_tag = _slugify(customer_phone, max_len=15)
    name_tag = _slugify(customer_name or fields.get("tenant_name", ""))
    filename = f"request_{request_id}_{phone_tag}_{name_tag}_{suffix}_{lang_tag}.docx"
    out_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(out_path)
    return out_path


def convert_to_pdf(docx_path: str) -> str:
    """
    Converts a .docx to .pdf, so the customer-facing watermarked draft can be
    sent as a WhatsApp-friendly PDF preview instead of a raw .docx.

    Tries headless LibreOffice first — the production-viable path, works the
    same on Mac (brew install --cask libreoffice) and Linux (apt-get install
    libreoffice). Falls back to docx2pdf (drives MS Word) for local Mac/
    Windows dev machines that have Word but not LibreOffice installed yet —
    NOT viable on a Linux production host, install LibreOffice there instead.
    """
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        outdir = os.path.dirname(docx_path)
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError("PDF conversion timed out after 60s")
        if result.returncode != 0:
            raise RuntimeError(f"PDF conversion failed: {result.stderr or result.stdout}")
        if not os.path.exists(pdf_path):
            raise RuntimeError(f"PDF conversion did not produce the expected file: {pdf_path}")
        return pdf_path

    # The docx2pdf/Word fallback drives MS Word via AppleScript, which has
    # proven unreliable in practice (crashes on some Word versions, or hangs
    # indefinitely waiting on a dialog Word popped up). Off by default so a
    # flaky local Word install can't hang approve requests; set
    # ENABLE_WORD_PDF_FALLBACK=true once you've confirmed it's stable on
    # your machine, or just install LibreOffice instead (recommended).
    if os.environ.get("ENABLE_WORD_PDF_FALLBACK", "").lower() != "true":
        raise RuntimeError(
            "No reliable PDF converter available. Install LibreOffice "
            "('brew install --cask libreoffice' locally, 'apt-get install "
            "libreoffice' on the production host). The MS Word fallback is "
            "disabled by default (unreliable) — set "
            "ENABLE_WORD_PDF_FALLBACK=true in .env to re-enable it."
        )

    try:
        from docx2pdf import convert
    except ImportError:
        raise RuntimeError(
            "No PDF converter available. Install LibreOffice ('brew install "
            "--cask libreoffice' locally, 'apt-get install libreoffice' on "
            "the production host) for the path that works everywhere."
        )
    try:
        convert(docx_path, pdf_path)
    except BaseException as e:
        # docx2pdf raises SystemExit (not a normal Exception) on failure,
        # e.g. AppleScript/Word automation errors — must be caught broadly
        # so callers relying on RuntimeError to gracefully fall back to the
        # .docx don't get a crash instead.
        raise RuntimeError(f"docx2pdf conversion via MS Word failed: {e}")
    if not os.path.exists(pdf_path):
        raise RuntimeError(f"PDF conversion did not produce the expected file: {pdf_path}")
    return pdf_path
